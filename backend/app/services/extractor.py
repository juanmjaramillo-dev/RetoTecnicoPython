import io

import docx
import openpyxl
import pdfplumber
from fastapi import UploadFile

ALLOWED_EXTENSIONS: set[str] = {".pdf", ".docx", ".xlsx"}


def _extract_pdf(data: bytes) -> str:
    parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []

    for para in document.paragraphs:
        line = para.text.strip()
        if line:
            parts.append(line)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _extract_xlsx(data: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    parts: list[str] = []

    for sheet in wb.worksheets:
        if sheet.title:
            parts.append(f"[Hoja: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip() not in ("", "None")]
            if not cells:
                continue
            # Si la fila tiene una sola celda larga, es probablemente un enunciado
            if len(cells) == 1:
                parts.append(cells[0])
            else:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text(file: UploadFile, data: bytes) -> str:
    filename = file.filename or ""
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""

    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Extensión no soportada: {ext!r}")

    if ext == ".pdf":
        return _extract_pdf(data)
    if ext == ".docx":
        return _extract_docx(data)
    return _extract_xlsx(data)
