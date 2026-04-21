"""
Script para generar el documento de muestra utilizado como ejemplo en /examples.
Requiere: pip install python-docx

Uso:
    cd examples
    python generate_sample.py
"""

import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_question_header(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f"{number}. {text}")
    run.bold = True
    run.font.size = Pt(11)


def build_document() -> Document:
    doc = Document()

    # Title
    title = doc.add_heading("Evaluación — Psicología Clínica: Trastornos de Ansiedad", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Instrucciones: lea atentamente cada enunciado y seleccione la alternativa correcta. "
        "En las preguntas de desarrollo, justifique su respuesta."
    )
    doc.add_paragraph()

    # Question 1 — selección múltiple (answer marked with asterisk)
    add_question_header(
        doc,
        1,
        "¿Cuál es el principal criterio diagnóstico del Trastorno de Ansiedad Generalizada (TAG)?",
    )
    doc.add_paragraph("A) Alucinaciones visuales recurrentes")
    doc.add_paragraph("B) Preocupación excesiva y difícil de controlar durante más de 6 meses  *")
    doc.add_paragraph("C) Estado de ánimo persistentemente elevado o expansivo")
    doc.add_paragraph("D) Episodios disociativos de amnesia")

    # Question 2 — verdadero/falso (answer stated explicitly)
    add_question_header(
        doc,
        2,
        "La terapia cognitivo-conductual (TCC) es el tratamiento psicológico "
        "con mayor evidencia empírica para el TAG.",
    )
    doc.add_paragraph("Verdadero / Falso")
    doc.add_paragraph("Respuesta correcta: Verdadero")

    # Question 3 — selección múltiple (answer marked with asterisk)
    add_question_header(
        doc,
        3,
        "¿Qué sistema de neurotransmisores está principalmente implicado en la fisiopatología del TAG?",
    )
    doc.add_paragraph("A) Dopaminérgico")
    doc.add_paragraph("B) Colinérgico")
    doc.add_paragraph("C) GABAérgico y serotoninérgico  *")
    doc.add_paragraph("D) Glutamatérgico")

    # Question 4 — verdadero/falso
    add_question_header(
        doc,
        4,
        "El Trastorno de Pánico puede presentarse con o sin agorafobia.",
    )
    doc.add_paragraph("Verdadero / Falso")
    doc.add_paragraph("Respuesta correcta: Verdadero")

    # Question 5 — selección múltiple
    add_question_header(
        doc,
        5,
        "¿Cuál de los siguientes fármacos es de primera línea en el tratamiento farmacológico del TAG?",
    )
    doc.add_paragraph("A) Haloperidol")
    doc.add_paragraph("B) Sertralina  *")
    doc.add_paragraph("C) Litio")
    doc.add_paragraph("D) Clozapina")

    # Question 6 — verdadero/falso
    add_question_header(
        doc,
        6,
        "La exposición gradual es una técnica exclusiva del tratamiento de fobias específicas "
        "y no se utiliza en otros trastornos de ansiedad.",
    )
    doc.add_paragraph("Verdadero / Falso")
    doc.add_paragraph("Respuesta correcta: Falso")

    # Question 7 — selección múltiple
    add_question_header(
        doc,
        7,
        "Según el DSM-5, ¿cuánto tiempo deben persistir los síntomas del TAG para cumplir criterio diagnóstico?",
    )
    doc.add_paragraph("A) Al menos 2 semanas")
    doc.add_paragraph("B) Al menos 1 mes")
    doc.add_paragraph("C) Al menos 3 meses")
    doc.add_paragraph("D) Al menos 6 meses  *")

    # Question 8 — verdadero/falso
    add_question_header(
        doc,
        8,
        "El modelo cognitivo de Beck propone que la ansiedad surge de la "
        "sobreestimación del peligro y la subestimación de los recursos de afrontamiento.",
    )
    doc.add_paragraph("Verdadero / Falso")
    doc.add_paragraph("Respuesta correcta: Verdadero")

    # Question 9 — selección múltiple
    add_question_header(
        doc,
        9,
        "¿Cuál de las siguientes NO es una técnica de relajación utilizada en el "
        "tratamiento de los trastornos de ansiedad?",
    )
    doc.add_paragraph("A) Respiración diafragmática")
    doc.add_paragraph("B) Relajación muscular progresiva de Jacobson")
    doc.add_paragraph("C) Entrenamiento autógeno de Schultz")
    doc.add_paragraph("D) Terapia electroconvulsiva  *")

    # Question 10 — verdadero/falso
    add_question_header(
        doc,
        10,
        "La fobia social y el trastorno de ansiedad generalizada son la misma entidad diagnóstica.",
    )
    doc.add_paragraph("Verdadero / Falso")
    doc.add_paragraph("Respuesta correcta: Falso")

    # Question 11 — selección múltiple
    add_question_header(
        doc,
        11,
        "¿Qué característica diferencia al Trastorno Obsesivo-Compulsivo (TOC) "
        "de los trastornos de ansiedad según el DSM-5?",
    )
    doc.add_paragraph("A) Presencia de pensamientos intrusivos")
    doc.add_paragraph("B) El TOC se clasifica en un capítulo diagnóstico separado  *")
    doc.add_paragraph("C) La ansiedad no aparece en el TOC")
    doc.add_paragraph("D) El TOC no responde a TCC")

    # Question 12 — verdadero/falso
    add_question_header(
        doc,
        12,
        "Las benzodiazepinas son recomendadas como tratamiento de mantenimiento "
        "a largo plazo para el TAG por su alto perfil de seguridad.",
    )
    doc.add_paragraph("Verdadero / Falso")
    doc.add_paragraph("Respuesta correcta: Falso")

    # Question 13 — desarrollo (open-ended, no answer)
    add_question_header(
        doc,
        13,
        "Describe brevemente la diferencia entre ansiedad adaptativa y ansiedad patológica, "
        "e indica al menos dos variables que determinan dicha distinción.",
    )
    doc.add_paragraph("Respuesta:")
    doc.add_paragraph("_" * 80)
    doc.add_paragraph("_" * 80)

    return doc


if __name__ == "__main__":
    document = build_document()
    output_path = os.path.join(os.path.dirname(__file__), "muestra_preguntas.docx")
    document.save(output_path)
    print(f"Archivo generado: {output_path}")
